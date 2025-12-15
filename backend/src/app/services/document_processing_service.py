"""
Document Processing Service - Parse, extract, and process documents.

WHY:
- Extract text from various file formats (PDF, DOCX, TXT, etc.)
- Handle file uploads and URL crawling
- Prepare documents for chunking and embedding
- Support metadata extraction

HOW:
- Use format-specific parsers (PyPDF2, docx, etc.)
- Extract text and metadata
- Clean and normalize text
- Handle errors gracefully

PSEUDOCODE follows the existing codebase patterns.
"""

import os
from typing import Optional
from uuid import UUID
import tempfile

from sqlalchemy.orm import Session


class DocumentProcessingService:
    """
    Document parsing and text extraction.

    WHY: Convert documents to text for embedding
    HOW: Format-specific parsers with error handling
    """

    def process_file(
        self,
        db: Session,
        file_path: str,
        kb_id: UUID,
        document_name: str,
        metadata: Optional[dict] = None
    ) -> dict:
        """
        Process uploaded file.

        WHY: Extract text from file for embedding
        HOW: Detect format, parse, extract text

        ARGS:
            db: Database session
            file_path: Path to uploaded file
            kb_id: Knowledge base ID
            document_name: Display name
            metadata: Additional metadata

        RETURNS:
            {
                "document_id": "uuid",
                "text": "Extracted text...",
                "pages": 10,
                "metadata": {...}
            }
        """

        # Detect file type
        file_ext = os.path.splitext(file_path)[1].lower()

        # Extract text based on format
        if file_ext == ".pdf":
            result = self._process_pdf(file_path)
        elif file_ext in [".docx", ".doc"]:
            result = self._process_docx(file_path)
        elif file_ext == ".txt":
            result = self._process_txt(file_path)
        elif file_ext in [".html", ".htm"]:
            result = self._process_html(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Create document record
        from app.models.document import Document

        document = Document(
            workspace_id=None,  # Will be set from KB
            kb_id=kb_id,
            name=document_name,
            source_type="file",
            source=file_path,
            content=result["text"],
            metadata={
                **(metadata or {}),
                "pages": result.get("pages"),
                "format": file_ext,
                "size_bytes": os.path.getsize(file_path)
            },
            status="processing"
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        return {
            "document_id": str(document.id),
            "text": result["text"],
            "pages": result.get("pages"),
            "metadata": document.metadata
        }


    def _process_pdf(self, file_path: str) -> dict:
        """
        Extract text from PDF.

        WHY: PDF is common document format
        HOW: Use PyPDF2 or pdfplumber
        """

        try:
            import PyPDF2

            text = ""
            pages = 0

            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                pages = len(reader.pages)

                for page in reader.pages:
                    text += page.extract_text() + "\n\n"

            return {
                "text": text.strip(),
                "pages": pages
            }

        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")


    def _process_docx(self, file_path: str) -> dict:
        """
        Extract text from DOCX.

        WHY: Word documents are common
        HOW: Use python-docx library
        """

        try:
            from docx import Document

            doc = Document(file_path)
            text = "\n\n".join([para.text for para in doc.paragraphs])

            return {
                "text": text.strip(),
                "pages": None  # DOCX doesn't have fixed pages
            }

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")


    def _process_txt(self, file_path: str) -> dict:
        """
        Read plain text file.

        WHY: Simplest format
        HOW: Direct file read
        """

        try:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()

            return {
                "text": text.strip(),
                "pages": None
            }

        except Exception as e:
            raise ValueError(f"Failed to read TXT file: {e}")


    def _process_html(self, file_path: str) -> dict:
        """
        Extract text from HTML.

        WHY: Web content
        HOW: Parse HTML, extract text content
        """

        try:
            from bs4 import BeautifulSoup

            with open(file_path, "r", encoding="utf-8") as file:
                html = file.read()

            soup = BeautifulSoup(html, "html.parser")

            # Remove script and style tags
            for script in soup(["script", "style"]):
                script.decompose()

            text = soup.get_text()

            # Clean whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            return {
                "text": text.strip(),
                "pages": None
            }

        except Exception as e:
            raise ValueError(f"Failed to parse HTML: {e}")


    async def process_url(
        self,
        db: Session,
        url: str,
        kb_id: UUID,
        document_name: str,
        metadata: Optional[dict] = None
    ) -> dict:
        """
        Process URL (web scraping).

        WHY: Add web content to KB
        HOW: Fetch, parse HTML, extract text

        ARGS:
            db: Database session
            url: URL to scrape
            kb_id: Knowledge base ID
            document_name: Display name
            metadata: Additional metadata

        RETURNS:
            {
                "document_id": "uuid",
                "text": "Extracted text...",
                "metadata": {...}
            }
        """

        import requests

        try:
            # Fetch URL
            response = requests.get(url, timeout=30)
            response.raise_for_status()

            # Save to temp file
            with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False) as tmp:
                tmp.write(response.text)
                tmp_path = tmp.name

            # Process as HTML
            result = self._process_html(tmp_path)

            # Clean up temp file
            os.unlink(tmp_path)

            # Create document record
            from app.models.document import Document

            document = Document(
                workspace_id=None,
                kb_id=kb_id,
                name=document_name,
                source_type="url",
                source=url,
                content=result["text"],
                metadata={
                    **(metadata or {}),
                    "url": url,
                    "content_type": response.headers.get("content-type")
                },
                status="processing"
            )

            db.add(document)
            db.commit()
            db.refresh(document)

            return {
                "document_id": str(document.id),
                "text": result["text"],
                "metadata": document.metadata
            }

        except Exception as e:
            raise ValueError(f"Failed to process URL: {e}")


    def process_text(
        self,
        db: Session,
        text: str,
        kb_id: UUID,
        document_name: str,
        metadata: Optional[dict] = None
    ) -> dict:
        """
        Process raw text input.

        WHY: Direct text input
        HOW: Create document directly

        ARGS:
            db: Database session
            text: Raw text content
            kb_id: Knowledge base ID
            document_name: Display name
            metadata: Additional metadata

        RETURNS:
            {
                "document_id": "uuid",
                "text": "Text content...",
                "metadata": {...}
            }
        """

        from app.models.document import Document

        document = Document(
            workspace_id=None,
            kb_id=kb_id,
            name=document_name,
            source_type="text",
            source="direct_input",
            content=text,
            metadata=metadata or {},
            status="processing"
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        return {
            "document_id": str(document.id),
            "text": text,
            "metadata": document.metadata
        }


# Global instance
document_processing_service = DocumentProcessingService()
